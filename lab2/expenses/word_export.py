"""Builds the .docx (editable Word) expense report from an ExpenseReport.

This is the second permanent record generated at approval time, alongside
the .xlsx from expenses/excel.py — see expenses/services.py:finalize_approval.
It captures the same structured data as the Excel export, plus an embedded
capture of *every* receipt — PDF invoices included, rendered to an image via
receipt_capture.py (PyMuPDF), not just photo receipts (JPG/PNG) — ordered by
expense date, so the captures read the same order as the trip happened in.

This function MUST run before the original TravelDocument.file is deleted —
once that happens there's nothing left to capture an image from.
"""
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt

from .policies import DAILY_LIMIT_USD, USD_MXN_RATE
from .receipt_capture import render_receipt_thumbnail

# Same green the company's real Advance & Expense Report spreadsheet uses
# to highlight its own totals column. VIOLATION_SHADE matches excel.py's
# VIOLATION_FILL so a flagged day reads the same color in both formats.
USD_SHADE = "C6EFCE"
VIOLATION_SHADE = "F8D7DA"


def _shade_cell(cell, hex_color: str) -> None:
    """Sets a table cell's background color — python-docx has no shading
    API of its own, so this edits the cell's XML directly (the documented
    way to do it)."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shading)


def build_report_document(report) -> Document:
    document = Document()

    document.add_heading("Travel expense report", level=1)

    trip_range = report.trip_date_range
    trip_range_text = f"{trip_range[0]:%Y-%m-%d} to {trip_range[1]:%Y-%m-%d}" if trip_range else "—"

    info_table = document.add_table(rows=0, cols=2)
    info_table.style = "Light Grid Accent 1"
    for label, value in [
        ("Employee", report.user.get_full_name() or report.user.email),
        ("Employee #", report.user.employee_number or "—"),
        ("Department", report.user.department),
        ("Supervisor", f"{report.supervisor_name} ({report.supervisor_email})" if report.supervisor_email else report.supervisor_name),
        ("Title", report.title),
        ("Trip dates", trip_range_text),
        ("Status", report.get_status_display()),
        ("Created", report.created_at.strftime("%Y-%m-%d %H:%M")),
    ]:
        row = info_table.add_row().cells
        row[0].text, row[1].text = label, value

    if report.reviewed_at:
        for label, value in [
            ("Reviewed", report.reviewed_at.strftime("%Y-%m-%d %H:%M")),
            ("Review note", report.review_note or "—"),
        ]:
            row = info_table.add_row().cells
            row[0].text, row[1].text = label, value
        if report.approval_clause:
            row = info_table.add_row().cells
            row[0].text, row[1].text = "Approval clause", report.approval_clause

    # Column layout mirrors the company's real Advance & Expense Report
    # spreadsheet (Invoice date / Description / Vendor legal name / FX
    # Rate / ... ), with "Amount (USD)" replacing that sheet's "Total MXN"
    # column since this app's policy and totals are USD-based.
    document.add_heading("Expenses", level=2)
    headers = [
        "#", "Invoice date", "Description", "Vendor legal name",
        "Expensed amount\nin foreign currency", "Currency", "FX Rate",
        "Backup", "File", "Amount (USD)",
    ]
    expense_table = document.add_table(rows=1, cols=len(headers))
    expense_table.style = "Light Grid Accent 1"
    header_cells = expense_table.rows[0].cells
    for cell, text in zip(header_cells, headers):
        cell.text = text
        cell.paragraphs[0].runs[0].bold = True
        if text == "Amount (USD)":
            _shade_cell(cell, USD_SHADE)

    # Days over the $60/day policy (see ExpenseReport.daily_totals — the one
    # place that decides this, currency conversion and the flight/hotel
    # exemption included) are flagged directly on their row here instead of
    # in a separate table: a whole extra section repeating dates the reader
    # just saw would only add noise, not information.
    violation_dates = {day["date"] for day in report.daily_totals() if day["over_limit"]}

    total_usd = 0
    # Alphabetically by expense type, then chronologically within each type.
    for index, doc in enumerate(report.documents.order_by("type", "document_date"), start=1):
        is_violation_day = doc.document_date in violation_dates
        row = expense_table.add_row().cells
        row[0].text = str(index)
        row[1].text = f"{doc.document_date:%Y-%m-%d}" + (" ⚠" if is_violation_day else "")
        row[2].text = doc.get_type_display()
        row[3].text = doc.vendor_name or "—"
        row[4].text = f"{doc.amount:.2f}"
        row[5].text = doc.currency
        row[6].text = "1.00" if doc.currency == "USD" else f"{USD_MXN_RATE:.2f}"
        row[7].text = doc.get_backup_type_display()
        row[8].text = doc.file_name
        row[9].text = f"${doc.amount_usd:.2f}"
        _shade_cell(row[9], USD_SHADE)
        if is_violation_day:
            row[1].paragraphs[0].runs[0].bold = True
            _shade_cell(row[1], VIOLATION_SHADE)
        total_usd += doc.amount_usd

    for subtotal in report.totals_by_currency():
        subtotal_paragraph = document.add_paragraph()
        subtotal_paragraph.add_run(f"Total ({subtotal['currency']}): {subtotal['total']:.2f}").bold = True

    total_paragraph = document.add_paragraph()
    total_run = total_paragraph.add_run(f"Total (USD): ${total_usd:.2f}")
    total_run.bold = True
    total_run.font.size = Pt(12)

    if violation_dates:
        dates_text = ", ".join(day.strftime("%Y-%m-%d") for day in sorted(violation_dates))
        policy_paragraph = document.add_paragraph()
        policy_run = policy_paragraph.add_run(
            f"⚠ Day(s) marked above exceed the ${DAILY_LIMIT_USD}/day policy and aren't "
            f"explained by a flight or hotel charge: {dates_text}"
        )
        policy_run.bold = True

    _add_receipt_captures(document, report)

    return document


def _add_receipt_captures(document: Document, report) -> None:
    """Embeds a visual capture of every receipt still on disk — a PDF
    invoice's first page rendered to an image, or a photo receipt as-is
    (see receipt_capture.py) — ordered by expense date, not type, so the
    captures read in the order the trip actually happened. Called before
    the originals are deleted — this is the permanent visual record of
    what every receipt looked like, invoices included, not just photos."""
    docs_with_files = [doc for doc in report.documents.order_by("document_date") if doc.file]
    if not docs_with_files:
        return

    document.add_heading("Receipt captures", level=2)
    for doc in docs_with_files:
        caption = document.add_paragraph()
        caption.add_run(f"{doc.get_type_display()} — {doc.document_date:%Y-%m-%d} — {doc.amount:.2f} {doc.currency}").bold = True
        try:
            thumbnail = render_receipt_thumbnail(doc.file)
            document.add_picture(BytesIO(thumbnail), width=Inches(3))
            document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT
        except Exception:  # noqa: BLE001 - a bad/unreadable receipt must never block the export
            document.add_paragraph("(Could not capture this receipt.)")
