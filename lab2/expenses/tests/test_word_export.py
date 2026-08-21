import base64
import shutil
import tempfile
from io import BytesIO

from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase, override_settings
from docx import Document

from accounts.models import User
from expenses.models import ExpenseReport, TravelDocument
from expenses.tests.helpers import make_pdf_bytes
from expenses.word_export import build_report_document

MEDIA_ROOT = tempfile.mkdtemp(prefix="expense_reports_tests_")

# A real (if tiny) 1x1 transparent PNG — needed because python-docx actually
# decodes the image to size it when embedding, so arbitrary bytes won't do.
ONE_PIXEL_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
)


@override_settings(MEDIA_ROOT=MEDIA_ROOT)
class WordExportTests(TestCase):
    @classmethod
    def tearDownClass(cls):
        super().tearDownClass()
        shutil.rmtree(MEDIA_ROOT, ignore_errors=True)

    def setUp(self):
        self.employee = User.objects.create_user(
            username="ana@example.com",
            email="ana@example.com",
            password="x",
            first_name="Ana Perez",
            department="Sales",
        )
        self.report = ExpenseReport.objects.create(
            user=self.employee, title="Trip to Mexico City", supervisor_name="Maria Lopez"
        )
        TravelDocument.objects.create(
            expense_report=self.report,
            file=SimpleUploadedFile("flight.pdf", b"x", content_type="application/pdf"),
            type=TravelDocument.DocType.FLIGHT,
            amount="2500.00",
            document_date="2026-08-01",
        )
        TravelDocument.objects.create(
            expense_report=self.report,
            file=SimpleUploadedFile("hotel.jpg", ONE_PIXEL_PNG, content_type="image/jpeg"),
            type=TravelDocument.DocType.HOTEL,
            amount="1800.00",
            document_date="2026-08-02",
        )

    def _all_text(self, document: Document) -> str:
        parts = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)

    def test_includes_employee_and_expense_data(self):
        document = build_report_document(self.report)
        text = self._all_text(document)

        self.assertIn("Ana Perez", text)
        self.assertIn("Sales", text)
        self.assertIn("Maria Lopez", text)
        self.assertIn("flight.pdf", text)
        self.assertIn("hotel.jpg", text)
        self.assertIn("$4300.00", text)

    def test_document_is_saveable_and_reopenable(self):
        document = build_report_document(self.report)
        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)

        reopened = Document(buffer)
        self.assertGreater(len(reopened.paragraphs) + len(reopened.tables), 0)

    def test_embeds_a_thumbnail_for_photo_receipts(self):
        document = build_report_document(self.report)
        # python-docx exposes embedded images via inline_shapes.
        self.assertGreaterEqual(len(document.inline_shapes), 1)

    def test_embeds_a_capture_for_pdf_receipts_too(self):
        # Not just photos — a real PDF invoice must also get a rendered
        # page image, not just be listed by name.
        report = ExpenseReport.objects.create(user=self.employee, title="PDF-only trip")
        TravelDocument.objects.create(
            expense_report=report,
            file=SimpleUploadedFile("invoice.pdf", make_pdf_bytes("Invoice"), content_type="application/pdf"),
            type=TravelDocument.DocType.TAXI,
            amount="50.00",
            document_date="2026-08-01",
        )

        document = build_report_document(report)

        self.assertGreaterEqual(len(document.inline_shapes), 1)
        text = self._all_text(document)
        self.assertNotIn("Could not capture this receipt", text)

    def test_no_documents_still_produces_a_document(self):
        empty_report = ExpenseReport.objects.create(user=self.employee, title="No documents")
        document = build_report_document(empty_report)
        self.assertGreaterEqual(len(document.paragraphs), 1)
