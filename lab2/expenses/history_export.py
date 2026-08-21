"""Builds the .docx "report history" download — the report's full audit
trail (every status change, who made it, and any review note, rejection
notes included) from creation through to its current state.

Unlike excel.py/word_export.py this is never a frozen snapshot: it's
generated live every time it's requested (see the admin's "Download
history" object-tool link, expenses/admin/reports.py), because the whole
point is to see the trail as it stands right now — including while a
report is still going through one or more reject/resubmit cycles, well
before it reaches approval.
"""
from docx import Document
from docx.shared import Pt


def build_history_document(report) -> Document:
    document = Document()
    document.add_heading("Report history", level=1)

    info_table = document.add_table(rows=0, cols=2)
    info_table.style = "Light Grid Accent 1"
    for label, value in [
        ("Report", report.title),
        ("Employee", report.user.get_full_name() or report.user.email),
        ("Employee #", report.user.employee_number or "—"),
        ("Current status", report.get_status_display()),
    ]:
        row = info_table.add_row().cells
        row[0].text, row[1].text = label, value

    document.add_heading("Timeline", level=2)
    headers = ["Date", "By", "Action", "Note"]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        cell.paragraphs[0].runs[0].bold = True

    for entry in report.audit_log.order_by("created_at"):
        row = table.add_row().cells
        row[0].text = entry.created_at.strftime("%Y-%m-%d %H:%M")
        row[1].text = entry.actor.get_full_name() or entry.actor.email if entry.actor else "System"
        row[2].text = entry.get_action_display()
        row[3].text = entry.note or "—"

    footer = document.add_paragraph()
    footer_run = footer.add_run(
        "This reflects every event recorded so far — including any rejection notes from "
        "earlier review cycles — not just the most recent decision."
    )
    footer_run.italic = True
    footer_run.font.size = Pt(9)

    return document
